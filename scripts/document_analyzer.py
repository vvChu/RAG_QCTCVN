"""
Document Analyzer - Intelligent Document Classification

Automatically detects document types and extracts metadata:
- Detects if TT-BXD/ND-CP promulgates QCVN/TCVN
- Extracts standard numbers, titles, and issuing information
- Generates correct filenames and categorization
"""
import re
from pathlib import Path
from typing import Dict, Optional, List
import pdfplumber
from docx import Document

# Regex patterns for detection
PATTERNS = {
    'qcvn_promulgation': r'[Bb]an\s+hành.*?QCVN\s*(\d+):(\d{4})/BXD',
    'tcvn_promulgation': r'[Bb]an\s+hành.*?TCVN\s*(\d+)[:\-](\d{4})',
    'qcvn_standalone': r'QCVN\s*(\d+):(\d{4})/BXD',
    'tcvn_standalone': r'TCVN\s*(\d+)[:\-](\d{4})',
    'tt_bxd_number': r'TT-BXD.*?(\d+)/(\d{4})',
    'nd_cp_number': r'ND-CP.*?(\d+)/(\d{4})',
    'qcvn_title': r'QCVN\s*\d+:\d{4}/BXD\s+(.+?)(?:\n|National)',
    'tcvn_title': r'TCVN\s*\d+[:\-]\d{4}\s+(.+?)(?:\n|$)'
}

def extract_text_from_pdf(file_path: Path, max_pages: int = 5) -> str:
    """Extract text from first N pages of PDF using pdfplumber."""
    try:
        text = ""
        with pdfplumber.open(str(file_path)) as pdf:
            for i in range(min(max_pages, len(pdf.pages))):
                page_text = pdf.pages[i].extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
        return ""

def extract_text_from_docx(file_path: Path, max_paragraphs: int = 30) -> str:
    """Extract text from first N paragraphs of DOCX."""
    try:
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs[:max_paragraphs])
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
        return ""

def clean_title(title: str, full_text: str, match_end_pos: int) -> str:
    """Clean extracted title and handle multi-line/generic titles."""
    if not title:
        return ""
    
    title = title.strip()
    
    # If title is generic header, look for next line
    generic_headers = [
        "QUY CHUẨN KỸ THUẬT QUỐC GIA",
        "TIÊU CHUẨN QUỐC GIA",
        "NATIONAL STANDARDS",
        "NATIONALSTANDARDS"
    ]
    
    if title.upper() in generic_headers:
        print(f"DEBUG: Generic header detected: '{title}'")
        # Look for next line in full_text starting from match_end_pos
        remaining_text = full_text[match_end_pos:]
        # Find first non-empty line
        lines = remaining_text.splitlines()
        for line in lines:
            line = line.strip()
            if line:
                print(f"DEBUG: Found next line: '{line}'")
                return line
        
        print("DEBUG: No next line found, returning generic title")
            
    return title

def detect_promulgated_standard(text: str) -> Dict:
    """
    Detect if document promulgates a QCVN/TCVN standard.
    """
    result = {
        'has_standard': False,
        'standard_type': None,
        'standard_number': None,
        'standard_title': None,
        'issuing_doc_type': None,
        'issuing_doc_number': None
    }
    
    # Check for QCVN promulgation
    qcvn_prom = re.search(PATTERNS['qcvn_promulgation'], text, re.IGNORECASE)
    if qcvn_prom:
        result['has_standard'] = True
        result['standard_type'] = 'QCVN'
        result['standard_number'] = f"{qcvn_prom.group(1)}:{qcvn_prom.group(2)}"
        
        # Extract title
        title_match = re.search(PATTERNS['qcvn_title'], text, re.IGNORECASE)
        if title_match:
            result['standard_title'] = clean_title(title_match.group(1), text, title_match.end())
        
        # Detect issuing document
        tt_bxd = re.search(PATTERNS['tt_bxd_number'], text)
        nd_cp = re.search(PATTERNS['nd_cp_number'], text)
        
        if tt_bxd:
            result['issuing_doc_type'] = 'TT-BXD'
            result['issuing_doc_number'] = f"{tt_bxd.group(1)}/{tt_bxd.group(2)}"
        elif nd_cp:
            result['issuing_doc_type'] = 'ND-CP'
            result['issuing_doc_number'] = f"{nd_cp.group(1)}/{nd_cp.group(2)}"
        
        return result
    
    # Check for TCVN promulgation
    tcvn_prom = re.search(PATTERNS['tcvn_promulgation'], text, re.IGNORECASE)
    if tcvn_prom:
        result['has_standard'] = True
        result['standard_type'] = 'TCVN'
        result['standard_number'] = f"{tcvn_prom.group(1)}-{tcvn_prom.group(2)}"
        
        # Extract title
        title_match = re.search(PATTERNS['tcvn_title'], text)
        if title_match:
            result['standard_title'] = clean_title(title_match.group(1), text, title_match.end())
        
        # Detect issuing document
        tt_bxd = re.search(PATTERNS['tt_bxd_number'], text)
        nd_cp = re.search(PATTERNS['nd_cp_number'], text)
        
        if tt_bxd:
            result['issuing_doc_type'] = 'TT-BXD'
            result['issuing_doc_number'] = f"{tt_bxd.group(1)}/{tt_bxd.group(2)}"
        elif nd_cp:
            result['issuing_doc_type'] = 'ND-CP'
            result['issuing_doc_number'] = f"{nd_cp.group(1)}/{nd_cp.group(2)}"
        
        return result
    
    # Check for standalone QCVN
    qcvn_standalone = re.search(PATTERNS['qcvn_standalone'], text, re.IGNORECASE)
    if qcvn_standalone:
        result['has_standard'] = True
        result['standard_type'] = 'QCVN'
        result['standard_number'] = f"{qcvn_standalone.group(1)}:{qcvn_standalone.group(2)}"
        
        # Extract title
        title_match = re.search(PATTERNS['qcvn_title'], text, re.IGNORECASE)
        if title_match:
            result['standard_title'] = clean_title(title_match.group(1), text, title_match.end())
        
        return result
    
    # Check for standalone TCVN
    tcvn_standalone = re.search(PATTERNS['tcvn_standalone'], text, re.IGNORECASE)
    if tcvn_standalone:
        result['has_standard'] = True
        result['standard_type'] = 'TCVN'
        result['standard_number'] = f"{tcvn_standalone.group(1)}-{tcvn_standalone.group(2)}"
        
        # Extract title
        title_match = re.search(PATTERNS['tcvn_title'], text, re.IGNORECASE)
        if title_match:
            result['standard_title'] = clean_title(title_match.group(1), text, title_match.end())

def extract_text_from_docx(file_path: Path, max_paragraphs: int = 30) -> str:
    """Extract text from first N paragraphs of DOCX."""
    try:
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs[:max_paragraphs])
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
        return ""

def detect_promulgated_standard(text: str) -> Dict:
    """
    Detect if document promulgates a QCVN/TCVN standard.
    
    Returns:
        {
            'has_standard': bool,
            'standard_type': 'QCVN' | 'TCVN' | None,
            'standard_number': 'XX:YYYY' or 'XXXX-YYYY',
            'standard_title': str,
            'issuing_doc_type': 'TT-BXD' | 'ND-CP' | None,
            'issuing_doc_number': 'XX/YYYY'
        }
    """
    result = {
        'has_standard': False,
        'standard_type': None,
        'standard_number': None,
        'standard_title': None,
        'issuing_doc_type': None,
        'issuing_doc_number': None
    }
    
    # Check for QCVN promulgation
    qcvn_prom = re.search(PATTERNS['qcvn_promulgation'], text, re.IGNORECASE)
    if qcvn_prom:
        result['has_standard'] = True
        result['standard_type'] = 'QCVN'
        result['standard_number'] = f"{qcvn_prom.group(1)}:{qcvn_prom.group(2)}"
        
        # Extract title
        title_match = re.search(PATTERNS['qcvn_title'], text, re.IGNORECASE)
        if title_match:
            result['standard_title'] = title_match.group(1).strip()
        
        # Detect issuing document
        tt_bxd = re.search(PATTERNS['tt_bxd_number'], text)
        nd_cp = re.search(PATTERNS['nd_cp_number'], text)
        
        if tt_bxd:
            result['issuing_doc_type'] = 'TT-BXD'
            result['issuing_doc_number'] = f"{tt_bxd.group(1)}/{tt_bxd.group(2)}"
        elif nd_cp:
            result['issuing_doc_type'] = 'ND-CP'
            result['issuing_doc_number'] = f"{nd_cp.group(1)}/{nd_cp.group(2)}"
        
        return result
    
    # Check for TCVN promulgation
    tcvn_prom = re.search(PATTERNS['tcvn_promulgation'], text, re.IGNORECASE)
    if tcvn_prom:
        result['has_standard'] = True
        result['standard_type'] = 'TCVN'
        result['standard_number'] = f"{tcvn_prom.group(1)}-{tcvn_prom.group(2)}"
        
        # Extract title
        title_match = re.search(PATTERNS['tcvn_title'], text)
        if title_match:
            result['standard_title'] = title_match.group(1).strip()
        
        # Detect issuing document
        tt_bxd = re.search(PATTERNS['tt_bxd_number'], text)
        nd_cp = re.search(PATTERNS['nd_cp_number'], text)
        
        if tt_bxd:
            result['issuing_doc_type'] = 'TT-BXD'
            result['issuing_doc_number'] = f"{tt_bxd.group(1)}/{tt_bxd.group(2)}"
        elif nd_cp:
            result['issuing_doc_type'] = 'ND-CP'
            result['issuing_doc_number'] = f"{nd_cp.group(1)}/{nd_cp.group(2)}"
        
        return result
    
    # Check for standalone QCVN
    qcvn_standalone = re.search(PATTERNS['qcvn_standalone'], text)
    if qcvn_standalone:
        result['has_standard'] = True
        result['standard_type'] = 'QCVN'
        result['standard_number'] = f"{qcvn_standalone.group(1)}:{qcvn_standalone.group(2)}"
        
        # Extract title
        title_match = re.search(PATTERNS['qcvn_title'], text)
        if title_match:
            result['standard_title'] = title_match.group(1).strip()
        
        return result
    
    # Check for standalone TCVN
    tcvn_standalone = re.search(PATTERNS['tcvn_standalone'], text)
    if tcvn_standalone:
        result['has_standard'] = True
        result['standard_type'] = 'TCVN'
        result['standard_number'] = f"{tcvn_standalone.group(1)}-{tcvn_standalone.group(2)}"
        
        # Extract title
        title_match = re.search(PATTERNS['tcvn_title'], text)
        if title_match:
            result['standard_title'] = title_match.group(1).strip()
        
        return result
    
    return result

def analyze_document(file_path: Path) -> Dict:
    """
    Analyze a document and extract all metadata.
    
    Returns complete classification info.
    """
    # Extract text
    if file_path.suffix.lower() == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif file_path.suffix.lower() == '.docx':
        text = extract_text_from_docx(file_path)
    else:
        return {'error': f'Unsupported file type: {file_path.suffix}'}
    
    if not text:
        return {'error': 'Could not extract text from file'}
    
    # Detect standard
    metadata = detect_promulgated_standard(text)
    metadata['file_path'] = str(file_path)
    metadata['file_name'] = file_path.name
    metadata['file_ext'] = file_path.suffix
    
    return metadata

def main():
    """Test the analyzer."""
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python document_analyzer.py <file_path>")
        sys.exit(1)
    
    file_path = Path(sys.argv[1])
    if not file_path.exists():
        print(f"File not found: {file_path}")
        sys.exit(1)
    
    print(f"Analyzing: {file_path.name}")
    print("=" * 60)
    
    result = analyze_document(file_path)
    
    if 'error' in result:
        print(f"Error: {result['error']}")
    else:
        print(f"Has Standard: {result['has_standard']}")
        if result['has_standard']:
            print(f"Standard Type: {result['standard_type']}")
            print(f"Standard Number: {result['standard_number']}")
            print(f"Standard Title: {result['standard_title']}")
            if result['issuing_doc_type']:
                print(f"Issued by: {result['issuing_doc_type']} {result['issuing_doc_number']}")
        else:
            print("No standard detected (standalone document)")

if __name__ == "__main__":
    main()
